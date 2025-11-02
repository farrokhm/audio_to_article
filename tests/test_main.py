import os
from unittest.mock import MagicMock, patch
import pytest
from audio_to_article.main import (
    ensure_directories,
    transcribe_audio,
    draft_article,
    improve_article,
    export_to_docx,
    export_to_pdf,
)

# Fixture for a mock genai client
@pytest.fixture
def mock_genai_client():
    client = MagicMock()
    # Mock the generate_content method to return an object with a .text attribute
    generate_content_mock = MagicMock()
    generate_content_mock.text = "Mocked article text"
    client.models.generate_content.return_value = generate_content_mock
    return client

def test_ensure_directories(tmp_path):
    """Tests that ensure_directories creates the required subdirectories."""
    base_dir = tmp_path
    subdirs = ["test_audios", "test_transcripts"]
    ensure_directories(base_dir, subdirs=subdirs)
    for sub in subdirs:
        assert os.path.isdir(os.path.join(base_dir, sub))

@patch('whisper.load_model')
def test_transcribe_audio(mock_load_model):
    """Tests the transcribe_audio function by mocking the whisper model."""
    # Arrange
    mock_model = MagicMock()
    mock_model.transcribe.return_value = {"text": "This is a test transcript."}
    mock_load_model.return_value = mock_model
    
    # Act
    result = transcribe_audio("fake/path/to/audio.mp3")
    
    # Assert
    mock_load_model.assert_called_once_with("small")
    mock_model.transcribe.assert_called_once_with("fake/path/to/audio.mp3")
    assert result == "This is a test transcript."

def test_draft_article_success(mock_genai_client):
    """Tests successful article drafting."""
    # Arrange
    transcript = "This is a test transcript."
    
    # Act
    result = draft_article(
        transcript=transcript,
        foundation_model="gemini-test-flash",
        client=mock_genai_client,
    )
    
    # Assert
    assert result == "Mocked article text"
    mock_genai_client.models.generate_content.assert_called_once()

def test_draft_article_error(mock_genai_client, capsys):
    """Tests error handling in article drafting."""
    # Arrange
    mock_genai_client.models.generate_content.side_effect = Exception("API Error")
    
    # Act
    result = draft_article(
        transcript="some transcript",
        foundation_model="gemini-test-flash",
        client=mock_genai_client,
    )
    
    # Assert
    assert result is None
    captured = capsys.readouterr()
    assert "An error occurred during article generation: API Error" in captured.out

def test_improve_article(mock_genai_client):
    """Tests the improve_article function."""
    # Arrange
    article_text = "This is a draft."
    mock_genai_client.models.generate_content.return_value.text = "This is an improved draft."
    
    # Act
    result = improve_article(article_text, mock_genai_client)
    
    # Assert
    assert result == "This is an improved draft."
    mock_genai_client.models.generate_content.assert_called_once()

def test_export_to_docx(tmp_path):
    """Tests exporting content to a .docx file."""
    # Arrange
    test_text = "Hello World"
    file_path = tmp_path / "test.docx"
    
    # Act
    export_to_docx(test_text, str(file_path))
    
    # Assert
    assert file_path.exists()
    from docx import Document
    doc = Document(str(file_path))
    assert len(doc.paragraphs) > 0
    assert doc.paragraphs[0].text == "Hello World"

def test_export_to_pdf(tmp_path):
    """Tests exporting content to a .pdf file."""
    # Arrange
    test_text = "Hello PDF"
    file_path = tmp_path / "test.pdf"
    
    # Act
    export_to_pdf(test_text, str(file_path))
    
    # Assert
    assert file_path.exists()
    assert file_path.stat().st_size > 0
